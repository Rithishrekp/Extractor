import os
import logging

# ======================================================
# Suppress unwanted device / tokenizer logs
# ======================================================

# Force CPU (optional, since it's already default)
os.environ["CUDA_VISIBLE_DEVICES"] = "-1"

# Suppress HuggingFace / spaCy tokenizer warnings
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"
import warnings
warnings.filterwarnings("ignore")
# Suppress Thinc device announce like "Device set to use cpu"
os.environ["THINC_NO_DEVICE_ANNOUNCE"] = "1"

# Suppress spaCy info logs
logging.getLogger("spacy").setLevel(logging.ERROR)

# ======================================================
# Imports
# ======================================================
import re
import fitz  # PyMuPDF
import docx
import spacy
import pytesseract
import cv2

from PIL import Image
import numpy as np

# ✅ Change this path if installed somewhere else
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

from flask import Flask, request, render_template
from werkzeug.utils import secure_filename
from PIL import Image
# from keybert import KeyBERT
from deep_translator import GoogleTranslator

import nltk
#nltk.download('punkt')
for resource in ["punkt", "punkt_tab"]:
    try:
        nltk.data.find(f"tokenizers/{resource}")
    except LookupError:
        nltk.download(resource)



# --- NEW: summarization imports ---
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from transformers import pipeline
''' text = "This is a test. NLTK punkt tokenizer should now work. This sentence is also part of the test."
parser = PlaintextParser.from_string(text, Tokenizer("english"))
summarizer = LsaSummarizer()
summary = summarizer(parser.document, 2)

for sentence in summary:
    print(sentence) '''

# ======================================================
# Extended Government Document Entities Schema
# ======================================================
ENTITY_SCHEMA = {
    "IDENTITY": [
        "NAME", "AADHAAR", "PAN", "PASSPORT_NO", "VOTER_ID", "DL_NO",
        "DOB", "GENDER", "FATHER_NAME", "ADDRESS", "ID_NO"
    ],
    "ADDRESS": [
        "NAME", "ADDRESS", "PINCODE", "ACCOUNT_NO", "CONSUMER_NO",
        "RATION_NO", "PROPERTY_ID"
    ],
    "FINANCE": [
        "ACCOUNT_NO", "IFSC", "AMOUNT", "DATE", "ORG",
        "TRANSACTION_ID", "GSTIN", "TAX_AMOUNT", "UAN_NO"
    ],
    "EDUCATION": [
        "NAME", "DEGREE", "UNIVERSITY", "YEAR", "MARKS",
        "SCHOOL", "CLASS", "SCHOLARSHIP_NAME",
        "RESUME", "CURRICULUM VITAE", "CV", "BIO DATA", "EDUCATION", "ACADEMIC", "UNIVERSITY", "COLLEGE",
        "SCHOOL", "INSTITUTE", "ACADEMY", "FACULTY", "CAMPUS", "DEGREE", "BACHELOR", "MASTER", "PhD",
        "CGPA", "PERCENTAGE", "GRADE", "MARKS", "TRANSCRIPT", "COURSE", "PROGRAM", "SYLLABUS", "STUDENT",
        "ROLL NUMBER", "ADMISSION", "PASSOUT", "GRADUATION", "QUALIFICATION", "SKILLS", "CERTIFICATION",
        "TRAINING", "INTERNSHIP", "PROJECT", "ACHIEVEMENTS", "EXPERIENCE", "WORKSHOP", "SEMINAR",
        "PUBLICATION", "REFERENCE"
    ],

    "HEALTH": [
        "PATIENT_NAME", "DOCTOR", "HOSPITAL", "DIAGNOSIS", "DATE",
        "VACCINE", "DOSE_NO", "DISABILITY_TYPE", "PERCENTAGE",
        "POLICY_NO", "INSURER", "COVERAGE_AMOUNT"
    ],
    "EMPLOYMENT": [
        "NAME", "ORG", "DESIGNATION", "SALARY", "PERIOD", "DATE",
        "EMPLOYEE_ID", "DEPARTMENT", "LABOUR_ID", "CATEGORY"
    ],
    "LEGAL": [
        "NAME", "DOB", "DOD", "DATE", "PLACE", "FATHER_NAME",
        "MOTHER_NAME", "HUSBAND_NAME", "WIFE_NAME", "CASTE",
        "STATE", "INCOME", "COURT_NAME"
    ],
    "PROPERTY": [
        "OWNER_NAME", "PROPERTY_ID", "ADDRESS", "REG_DATE",
        "REG_OFFICE", "SURVEY_NO", "AREA", "LOCATION"
    ],
    "VEHICLE": [
        "RC_NO", "OWNER_NAME", "VEHICLE_NO", "MODEL", "DATE",
        "POLICY_NO", "INSURER", "PERMIT_NO", "VEHICLE_CLASS"
    ],
    "BUSINESS": [
        "LICENSE_NO", "ORG", "OWNER_NAME", "DATE", "ADDRESS",
        "IEC_NO", "UAM_NO", "COUNTRY"
    ],
    "WELFARE": [
        "CARD_NO", "JOB_CARD_NO", "PENSION_NO", "NAME", "DOB",
        "AMOUNT", "FAMILY_MEMBERS", "ADDRESS"
    ],
    "GENERAL": [
        "PERSON", "ORG", "DATE", "PLACE"
    ]
}

# ----------------------------
# Flask setup
# ----------------------------
app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# ----------------------------
# NLP / Keyword models
# ----------------------------
nlp = spacy.load("en_core_web_sm")
#kw_model = KeyBERT()

# --- NEW: abstractive model load (once) ---
summarizer_model = pipeline("summarization", model="facebook/bart-large-cnn")

# ----------------------------
# Allowed file types
# ----------------------------
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "png", "jpg", "jpeg"}

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# ----------------------------
# Helpers
# ----------------------------
''' def ocr_image(image_path, lang="eng"):
    # Read image
    img = cv2.imread(image_path)

    # Preprocess (grayscale + threshold)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    gray = cv2.threshold(gray, 0, 255,
                         cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

    # OCR extraction
    text = pytesseract.image_to_string(gray, lang=lang)
    return text '''

def clean_list(values):
    seen = set()
    out = []
    for v in values:
        v = v.strip()
        if not v:
            continue
        if v not in seen:
            seen.add(v)
            out.append(v)
    return out

def clean_entities(entities):
    cleaned = {}
    for label, items in entities.items():
        unique_items = set(items)  # remove duplicates
        filtered = []

        for item in unique_items:
            text = item.strip()
            # Skip unwanted misclassifications
            if label == "PERSON":
                if any(char.isdigit() for char in text):
                    continue
                if len(text.split()) > 4:
                    continue
            if label == "DATE":
                if len(text) < 3:
                    continue
            if label == "ORG":
                if len(text) < 3:
                    continue
            filtered.append(text)

        if filtered:
            cleaned[label] = filtered
    return cleaned

def is_phone_like(s: str) -> bool:
    s_compact = re.sub(r"[^\d+]", "", s)
    digits = re.sub(r"\D", "", s_compact)
    return 10 <= len(digits) <= 15

# ----------------------------
# Text extraction
# ----------------------------
''' def extract_text(filepath: str) -> str:
    text = ""
    ext = filepath.rsplit(".", 1)[1].lower()
    try:
        if ext == "pdf":
            doc = fitz.open(filepath)
            try:
                for page in doc:
                    text += page.get_text("text")
                if not text.strip():
                    for page_index in range(len(doc)):
                        pix = doc[page_index].get_pixmap(dpi=200)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        text += pytesseract.image_to_string(img) + "\n"
            finally:
                doc.close()
        elif ext == "docx":
            d = docx.Document(filepath)
            text = "\n".join([p.text for p in d.paragraphs])
        elif ext == "txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext in ["png", "jpg", "jpeg"]:
            img = cv2.imread(filepath)
            if img is None:
                return ""
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
            text = pytesseract.image_to_string(thresh)
    except Exception as e:
        text += f"\n[Extraction error: {e}]"
    return text '''
''' def extract_text(file_path):
    import cv2
    import pytesseract
    import numpy as np
    import os

    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".pdf":
        # your PDF extraction code here
        pass
    elif ext == ".docx":
        # your DOCX extraction code here
        pass
    elif ext == ".txt":
        # your TXT extraction code here
        pass
    elif ext in [".png", ".jpg", ".jpeg", ".tiff"]:
        # OCR with preprocessing
        img = cv2.imread(file_path)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        denoised = cv2.medianBlur(gray, 3)
        _, thresh = cv2.threshold(
            denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
        )
        kernel = np.array([[0, -1, 0],
                           [-1, 5, -1],
                           [0, -1, 0]])
        sharpened = cv2.filter2D(thresh, -1, kernel)

        text = pytesseract.image_to_string(
            sharpened, lang="eng", config="--oem 3 --psm 6"
        )

    return text'''
def extract_text(file_path):
    import cv2
    import pytesseract
    import numpy as np
    import os
    import fitz  # PyMuPDF for PDFs
    import docx  # python-docx for DOCX

    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".pdf":
        # Extract text from PDF
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text")
        doc.close()

    elif ext == ".docx":
        # Extract text from DOCX
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"

    elif ext == ".txt":
        # Extract text from TXT
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

    elif ext in [".png", ".jpg", ".jpeg", ".tiff"]:
        # OCR with preprocessing
        img = cv2.imread(file_path)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        denoised = cv2.medianBlur(gray, 3)
        _, thresh = cv2.threshold(
            denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
        )
        kernel = np.array([[0, -1, 0],
                           [-1, 5, -1],
                           [0, -1, 0]])
        sharpened = cv2.filter2D(thresh, -1, kernel)

        text = pytesseract.image_to_string(
            sharpened, lang="eng", config="--oem 3 --psm 6"
        )

    return text





# ----------------------------
# Regex extractors
# ----------------------------
PHONE_REGEXES = [
    r"(?:\+91[\s-]?)?[6-9]\d{9}\b",
    r"\b[6-9]\d{9}\b",

    r"\+\d{1,3}[\s-]?\d[\d\s-]{7,}\d",
]
EMAIL_REGEX = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"

MONEY_REGEX = re.compile(
    r'(?:(?:₹|Rs\.?|INR|USD|EUR|GBP|\$)\s?\d{1,3}(?:,\d{3})*(?:\.\d+)?'
    r'|\d+(?:,\d{3})*(?:\.\d+)?\s?(?:rupees|dollars|euros|pounds|inr|usd|eur|gbp|crore|crores|lakh|lakhs|million|billion|thousand))',
    re.IGNORECASE
)

PAN_REGEX = r"\b[A-Z]{5}[0-9]{4}[A-Z]\b"
AADHAAR_REGEX = r"\b\d{4}\s\d{4}\s\d{4}\b"
IFSC_REGEX = r"\b[A-Z]{4}0[A-Z0-9]{6}\b"
BANK_ACC_REGEX = r"\b\d{9,18}\b"
def filter_bank_accounts(candidates, phone_numbers):
    """Remove phone-like numbers from bank accounts"""
    cleaned = []
    for acc in candidates:
        # If already detected as phone, skip
        if any(acc.endswith(ph) or acc == ph for ph in phone_numbers):
            continue
        # If looks like phone with 91 prefix (12 digits starting with 91 + 10 digit)
        if acc.startswith("91") and len(acc) == 12:
            continue
        cleaned.append(acc)
    return cleaned

def extract_money(text: str):
    candidates = MONEY_REGEX.findall(text)
    flat = []
    for c in candidates:
        if isinstance(c, (list, tuple)):
            c = "".join(c)
        c = c.strip()
        if c:
            c = re.sub(r"(\d)\s(\d{2,3})", r"\1,\2", c)
            flat.append(c)
    return clean_list(flat)

DATE_REGEXES = [
    r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
    r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b",
    r"\b(?:\d{1,2}\s)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sept|Oct|Nov|Dec)[a-z]*\s?\d{2,4}\b",
    r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4}\b",
]

def extract_phones(text: str):
    phones = []
    for rx in PHONE_REGEXES:
        phones += re.findall(rx, text)
    return clean_list(phones)

def extract_emails(text: str):
    return clean_list(re.findall(EMAIL_REGEX, text))

def extract_dates(text: str, spacy_doc):
    dates = set()
    for ent in spacy_doc.ents:
        if ent.label_ == "DATE":
            s = ent.text.strip()
            if not is_phone_like(s):
                dates.add(s)
    for rx in DATE_REGEXES:
        for m in re.findall(rx, text, flags=re.IGNORECASE):
            s = m.strip()
            if not is_phone_like(s):
                dates.add(s)
    return clean_list(list(dates))

# ----------------------------
# Keyword Extraction
# ----------------------------
import re

# Invalid PERSON words that should not be in PERSON
invalid_person_keywords = {
    "regno", "registration", "account", "number", "no",
    "pan", "ifsc", "gst", "bachelor", "engineering", "science",
    "project", "course", "amount", "salary", "upi","street","nagar"
}

# Regex patterns for financial & IDs
patterns = {
    "PAN": r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
    "IFSC": r"\b[A-Z]{4}0[A-Z0-9]{6}\b",
    "ACCOUNT_NO": r"\b\d{9,18}\b"
}
address_keywords = {
    "street", "st", "st.", "road", "rd", "rd.", "nagar", "avenue",
    "colony", "layout", "phase", "block", "sector", "circle"
}


def extract_keywords(text: str):
    doc = nlp(text)
    out = {}
    # spaCy entities
    for ent in doc.ents:
        label = ent.label_
        val = ent.text.strip()
        if not val:
            continue
        out.setdefault(label, [])
        if val not in out[label]:
            out[label].append(val)
    phones = []
    for regex in PHONE_REGEXES:
        phones.extend(re.findall(regex, text))
    if phones:
        out["PHONE"] = list(set(phones))  # remove duplicates

    # Find bank accounts and filter out phone-like numbers
    bank_acc = re.findall(BANK_ACC_REGEX, text)
    if bank_acc:
        bank_acc = filter_bank_accounts(bank_acc, out.get("PHONE", []))
        if bank_acc:
            out["BANK_ACCOUNT"] = bank_acc

    # Regex-based
    if phones := extract_phones(text):
        out["PHONE"] = phones
    if emails := extract_emails(text):
        out["EMAIL"] = emails
    if money_list := extract_money(text):
        out.setdefault("MONEY", []).extend(money_list)
    if dates := extract_dates(text, doc):
        out.setdefault("DATE", []).extend(dates)
    if pan := re.findall(PAN_REGEX, text):
        out["PAN"] = pan
    if aadhaar := re.findall(AADHAAR_REGEX, text):
        out["AADHAAR"] = aadhaar
    if ifsc := re.findall(IFSC_REGEX, text):
        out["IFSC"] = ifsc
    if bank_acc := re.findall(BANK_ACC_REGEX, text):
        bank_acc = [acc for acc in bank_acc if not is_phone_like(acc)]
        if bank_acc:
            out["BANK_ACCOUNT"] = bank_acc

    # Domain-specific terms
    finance_terms = re.findall(r"(salary|loan|emi|interest|account|transaction|fund|invoice)", text, re.I)
    edu_terms = re.findall(r"(degree|university|certificate|marks|exam|semester)", text, re.I)
    health_terms = re.findall(r"(patient|doctor|hospital|diagnosis|treatment|prescription|blood pressure)", text, re.I)

    if finance_terms: out["FINANCE_TERMS"] = clean_list(finance_terms)
    if edu_terms: out["EDUCATION_TERMS"] = clean_list(edu_terms)
    if health_terms: out["HEALTH_TERMS"] = clean_list(health_terms)

    # Cleanup
    cleaned_out = {}
    for k, v in out.items():
        vals = clean_list(v)
        if k == "PERSON":
            vals = [x for x in vals if not any(y in x for y in ["Bhavan","Road","Gramin","DDWS","Department","Mission"]) and not any(char.isdigit() for char in x) and len(x.split())<=4]
        ''' if k == "ORG":
            vals = [x for x in vals if len(x.split())>1 or x.isupper()]'''
        if k == "ORG":
            new_vals = []
            for val in vals:
               if val.isupper() and len(val.split()) == 1:
            # Likely a PERSON, move to PERSON
                cleaned_out.setdefault("PERSON", []).append(val.title())
               else:
                new_vals.append(val)
            vals = new_vals

        if k == "DATE":
            vals = [x for x in vals if len(x)>=3 and not re.match(r"^\d{6}$",x) and not x.lower() in ["annually","annual"]]
        if k == "MONEY":
            vals = [re.sub(r"\s+","",x) for x in vals if not re.match(r"^#+",x)]
        if vals:
            cleaned_out[k]=vals
    return cleaned_out

# ----------------------------
# Document classification
# ----------------------------
# ----------------------------
# Document classification
# ----------------------------
def detect_document_field(text: str, keywords: dict) -> str:
    blob_parts = []
    for k, vals in keywords.items():
        if isinstance(vals, list):
            blob_parts.extend(vals)
    blob = (text + " " + " ".join(blob_parts)).lower()

    # Use regex word boundaries to avoid false positives (e.g., "accountability")
    financial_patterns = [r"\bsalary\b", r"\baccount\b", r"\bifsc\b", r"\bgst\b",
                          r"\bamount\b", r"\bfunding\b", r"\binvoice\b",
                          r"\btransaction\b", r"\bloan\b", r"\bbank\b"]

    education_patterns = [r"\bdegree\b", r"\buniversity\b", r"\binstitute\b",
                          r"\bschool\b", r"\bstudent\b", r"\bexam\b",
                          r"\bcertificate\b", r"\bmarks\b", r"\bsemester\b"]

    health_patterns = [r"\bhospital\b", r"\bpatient\b", r"\bdoctor\b",
                       r"\bprescription\b", r"\bmedical\b", r"\bhealth\b",
                       r"\btreatment\b", r"\bdiagnosis\b"]

    identity_patterns = [r"\baadhaar\b", r"\bpassport\b", r"\bvoter\b",
                         r"\blicense\b", r"\blicence\b", r"\bid card\b",
                         r"\bpan\b", r"\bdriving\b"]

    govt_patterns = [r"\bministry\b", r"\bgovernment\b", r"\bdepartment\b",
                     r"\bsecretary\b", r"\bgazette\b", r"\border\b",
                     r"\bnotification\b", r"\bscheme\b", r"\boffice memorandum\b",
                     r"\bcircular\b"]

    def has_any(patterns):
        return any(re.search(p, blob) for p in patterns)

    # --- PRIORITY FIX: Government docs should override financial ---
    if has_any(education_patterns):
        return "Education"
    if has_any(govt_patterns):
        return "Government"
    if has_any(financial_patterns):
        return "Financial"
    
    if has_any(health_patterns):
        return "Health "
    if has_any(identity_patterns):
        return "Identity"

    return "General"


# ----------------------------
# Summarization
# ----------------------------
def extractive_summary(text, sentence_count=3):
    clean_text = text.strip()
    if not clean_text:
        return "No content to summarize."
    try:
        parser = PlaintextParser.from_string(clean_text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        total_sentences = len(list(parser.document.sentences))
        actual_count = min(sentence_count, total_sentences)
        if actual_count==0:
            return "Text too short for extractive summary."
        summary_sentences = summarizer(parser.document, actual_count)
        summary_sentences = [str(s).strip() for s in summary_sentences if str(s).strip()]
        if not summary_sentences:
            return "Text too short or unsuitable for extractive summary."
        return " ".join(summary_sentences)
    except Exception as e:
        return f"Error in extractive summarization: {e}"

def abstractive_summary(text):
    if len(text.strip())==0:
        return "No content to summarize."
    text=text[:1024]
    summary = summarizer_model(text,max_length=120,min_length=30,do_sample=False)
    return summary[0]['summary_text']
from flask import Flask, request, jsonify, url_for
from datetime import datetime
import uuid
from gtts import gTTS
import os

# Generate speech

''' tts = gTTS("Hello, this is a test", lang="en")
os.makedirs("static", exist_ok=True)

# Save in static folder
BASE_DIR = os.path.dirname(os.path.abspath(_file_))
static_folder = os.path.join(BASE_DIR, "static")
os.makedirs(static_folder, exist_ok=True)
output_path = os.path.join("static", "voice_output.mp3")
tts.save(output_path)

print("File saved at:", output_path)


def speak_text(text, lang="en"):
    if not text:
        return None

    BASE_DIR = os.path.dirname(os.path.abspath(_file_))
    static_folder = os.path.join(BASE_DIR, "static")
    os.makedirs(static_folder, exist_ok=True)

    filename = os.path.join(static_folder, "voice_output.mp3")
    tts = gTTS(text=text, lang=lang)
    tts.save(filename)

    return "voice_output.mp3" '''
def summarize_text(text):
    # Use extractive summary by default
    return extractive_summary(text)

def translate_text(text, target_lang="en"):
    try:
        return GoogleTranslator(source="auto", target=target_lang).translate(text)
    except Exception as e:
        return f"[Translation error: {e}]"

'''def speak_text_dynamic(text, lang="en", prefix="audio"):
    try:
        BASE_DIR = os.path.dirname(os.path.abspath(_file_))
        static_folder = os.path.join(BASE_DIR, "static")
        os.makedirs(static_folder, exist_ok=True)

        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        filename = f"{prefix}{timestamp}{unique_id}.mp3"
        filepath = os.path.join(static_folder, filename)

        # Generate TTS
        tts = gTTS(text=text, lang=lang)
        tts.save(filepath)

        return filename  # return the filename for template playback
    except Exception as e:
        print("Error generating TTS:", e)
        return None '''
from gtts import gTTS
import os, uuid

def speak_text_dynamic(text, lang="en", prefix="audio"):
    if not text or not text.strip():
        print(">>> TTS INPUT is empty, skipping")
        return None

    audio_folder = os.path.join("static", "audio")
    os.makedirs(audio_folder, exist_ok=True)

    filename = f"{prefix}_{uuid.uuid4().hex}.mp3"
    file_path = os.path.join(audio_folder, filename)

    tts = gTTS(text=text.strip(), lang=lang)
    tts.save(file_path)

    return f"audio/{filename}"


    


# Generate audio


''' @app.route("/process", methods=["POST"])
def process_text():
    text = request.form.get("text")

    # Your existing processing
    summary = extractive_summary(text)   # or abstractive_summary
    translation = GoogleTranslator(source="auto", target="en").translate(summary)
    keywords = extract_keywords(text)
    field = detect_document_field(text, keywords)

    # --- Combine everything into one text for audio ---
    audio_text = f"""
    Detected Field: {field}.
    Summary: {summary}.
    Translation: {translation}.
    Extracted Keywords: {', '.join([f'{k}: {', '.join(v)}' for k, v in keywords.items()])}.
    """

    # --- Save audio ---
    audio_file = speak_text_dynamic(audio_text)

    # --- Return with audio file ---
    return render_template(
        "index.html",
        field=field,
        summary=summary,
        translation=translation,
        keywords=keywords,
        audio_file=audio_file
    )'''
''' @app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "No file uploaded"

    file = request.files["file"]
    if file.filename == "":
        return "No selected file"

    if file:
        file_path = os.path.join("uploads", file.filename)
        file.save(file_path)

        # =====================================================
        # 👉 Processing
        # =====================================================
        extracted_text = extract_text(file_path)
        keywords = extract_keywords(extracted_text)
        detected_field = detect_document_field(extracted_text, keywords)
        summary = summarize_text(extracted_text)
        translation = translate_text(extracted_text, target_lang="ta")
        full_text = extract_text(file_path)
        extracted = extract_keywords(full_text)
        detected_field = detect_document_field(full_text, extracted)
        ex_summary = summarize_text(full_text)
        translated_text = translate_text(full_text, target_lang="ta")
        
        


        # Generate individual audios
        audio_extracted = speak_text_dynamic(extracted_text, prefix="extracted")
        audio_keywords = speak_text_dynamic(", ".join(keywords), prefix="keywords")
        audio_summary = speak_text_dynamic(summary, prefix="summary")
        audio_translation = speak_text_dynamic(translation, prefix="translation") 
        target_lang = target_lang if 'target_lang' in locals() else "en"

        audio_extracted = speak_text_dynamic(full_text, prefix="extracted", lang=target_lang) if full_text else None

        audio_keywords = speak_text_dynamic(
           ", ".join([item for sublist in extracted.values() for item in sublist]),
           prefix="keywords",
           lang=target_lang
        ) if extracted else None

        audio_summary = speak_text_dynamic(ex_summary, prefix="summary", lang=target_lang) if ex_summary else None
        audio_translation = speak_text_dynamic(translated_text, prefix="translation", lang=target_lang) if translated_text else None


        # ✅ Generate combined narration
            # --- Combine everything for voice output ---
    full_text_for_voice = (
        f"Original text is: {full_text}. "
        f"Extracted details are: {extracted}. "
        f"Extractive summary is: {ex_summary}. "
        f"Abstractive summary is: {ab_summary}. "
        f"Translated text in {target_lang} is: {translated_text}. "
        f"Translated keywords are: {translated_keywords}. "
        f"Translated extractive summary is: {translated_ex_summary}. "
        f"Translated abstractive summary is: {translated_ab_summary}."
    )

    # --- Generate audio narration ---
    audio_combined = None
    if full_text_for_voice.strip():
        audio_combined = speak_text_dynamic(full_text_for_voice, lang="en", prefix="combined")

    # --- Render template with all results ---
    return render_template(
        "index.html",
        extracted=extracted,
        detected_field=detected_field,
        full_text=full_text,
        translated_text=translated_text,
        translated_keywords=translated_keywords,
        target_lang=target_lang,
        ex_summary=ex_summary,
        ab_summary=ab_summary,
        translated_ex_summary=translated_ex_summary,
        translated_ab_summary=translated_ab_summary,
        files_by_category=files_by_category,
        audio_combined=audio_combined
    ) 




# ----------------------------
# Routes with translation
# ----------------------------
#from flask import Flask, request, render_template, url_for
import os '''

''' @app.route("/", methods=["GET", "POST"])
def index():
    extracted = None
    detected_field = None
    full_text = ""
    translated_text = ""
    translated_keywords = {}
    target_lang = "en"
    ex_summary = ""
    ab_summary = ""
    translated_ex_summary = ""
    translated_ab_summary = ""

    # --- List already uploaded files ---
    files_by_category = {}
    for root, dirs, files in os.walk(app.config["UPLOAD_FOLDER"]):
        if files:
            category = os.path.basename(root) or "General"
            files_by_category[category] = [
                os.path.join(category, f) for f in files
            ]

    # --- Handle upload ---
    if request.method == "POST":
        if "document" not in request.files:
            return render_template("index.html", files_by_category=files_by_category)

        file = request.files["document"]
        if not file or file.filename == "":
            return render_template("index.html", files_by_category=files_by_category)

        if not allowed_file(file.filename):
            return render_template(
                "index.html",
                detected_field="Unsupported file type.",
                files_by_category=files_by_category
            )

        # Save file
        filename = secure_filename(file.filename)
        #filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        temp_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(temp_path)

        # --- Extract text & process ---
        full_text = extract_text(temp_path)
        extracted = extract_keywords(full_text)
        detected_field = detect_document_field(full_text, extracted)

        field_folder = os.path.join(app.config["UPLOAD_FOLDER"], detected_field)
        os.makedirs(field_folder, exist_ok=True)

    # Step 4: Move file into that folder
        new_path = os.path.join(field_folder, filename)
        os.replace(temp_path, new_path) 

        # Summarization
        if full_text.strip():
            try:
                ex_summary = extractive_summary(full_text)
            except Exception:
                ex_summary = "Error in extractive summarization."
            try:
                ab_summary = abstractive_summary(full_text)
            except Exception:
                ab_summary = "Error in abstractive summarization."

        # --- Translate ---
        target_lang = request.form.get("lang", "en")
        translated_text = full_text
        translated_keywords = extracted
        translated_ex_summary = ex_summary
        translated_ab_summary = ab_summary

        if target_lang != "en":
            try:
                translated_text = GoogleTranslator(source="auto", target=target_lang).translate(full_text)
                translated_keywords = {
                    k: [GoogleTranslator(source="auto", target=target_lang).translate(v) for v in vals]
                    for k, vals in extracted.items()
                }
                translated_ex_summary = GoogleTranslator(source="auto", target=target_lang).translate(ex_summary)
                translated_ab_summary = GoogleTranslator(source="auto", target=target_lang).translate(ab_summary)
            except Exception as e:
                translated_text = f"Translation error: {e}"
                translated_keywords = extracted
                translated_ex_summary = f"Translation error: {e}"
                translated_ab_summary = f"Translation error: {e}"

    # --- Render page ---
    return render_template(
        "index.html",
        extracted=extracted,
        detected_field=detected_field,
        full_text=full_text,
        translated_text=translated_text,
        translated_keywords=translated_keywords,
        target_lang=target_lang,
        ex_summary=ex_summary,
        ab_summary=ab_summary,
        translated_ex_summary=translated_ex_summary,
        translated_ab_summary=translated_ab_summary,
        files_by_category=files_by_category
    ) 
@app.route("/", methods=["GET", "POST"])
def index():
    extracted = None
    detected_field = None
    full_text = ""
    translated_text = ""
    translated_keywords = {}
    target_lang = "en"
    ex_summary = ""
    ab_summary = ""
    translated_ex_summary = ""
    translated_ab_summary = ""

    # --- Handle upload ---
    if request.method == "POST":
        if "document" in request.files:
            file = request.files["document"]
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                temp_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(temp_path)

                # --- Extract text & process ---
                full_text = extract_text(temp_path)
                extracted = extract_keywords(full_text)
                detected_field = detect_document_field(full_text, extracted)

                # Create field folder and move file
                field_folder = os.path.join(app.config["UPLOAD_FOLDER"], detected_field)
                os.makedirs(field_folder, exist_ok=True)
                new_path = os.path.join(field_folder, filename)
                os.replace(temp_path, new_path)

                # --- Summarization ---
                if full_text.strip():
                    try:
                        ex_summary = extractive_summary(full_text)
                    except Exception:
                        ex_summary = "Error in extractive summarization."
                    try:
                        ab_summary = abstractive_summary(full_text)
                    except Exception:
                        ab_summary = "Error in abstractive summarization."

                # --- Translation ---
                target_lang = request.form.get("lang", "en")
                translated_text = full_text
                translated_keywords = extracted
                translated_ex_summary = ex_summary
                translated_ab_summary = ab_summary

                if target_lang != "en":
                    try:
                        translated_text = GoogleTranslator(source="auto", target=target_lang).translate(full_text)
                        translated_keywords = {
                            k: [GoogleTranslator(source="auto", target=target_lang).translate(v) for v in vals]
                            for k, vals in extracted.items()
                        }
                        translated_ex_summary = GoogleTranslator(source="auto", target=target_lang).translate(ex_summary)
                        translated_ab_summary = GoogleTranslator(source="auto", target=target_lang).translate(ab_summary)
                    except Exception as e:
                        translated_text = f"Translation error: {e}"
                        translated_keywords = extracted
                        translated_ex_summary = f"Translation error: {e}"
                        translated_ab_summary = f"Translation error: {e}"

    # --- Always dynamically scan uploads folder ---
    files_by_category = {}
    uploads_dir = app.config["UPLOAD_FOLDER"]

    for root, dirs, filenames in os.walk(uploads_dir):
        rel_dir = os.path.relpath(root, uploads_dir)
        if rel_dir == ".":
            rel_dir = "Uncategorized"
        files_by_category.setdefault(rel_dir, [])
        for f in filenames:
            rel_path = os.path.relpath(os.path.join(root, f), uploads_dir)
            rel_path = rel_path.replace("\\", "/")
            files_by_category[rel_dir].append(rel_path)
    # --- Generate combined text for TTS ---
    # --- Generate combined text for TTS ---
    

        # --- Combine everything for voice output ---
    full_text_for_voice = f"""
    Original Text:
    {full_text}

    Extracted Details:
    {extracted}

   Extractive Summary:
   {ex_summary}

   Abstractive Summary:
   {ab_summary}

   Text ({target_lang}):
   {translated_text}

   Translated Extracted Details:
   {translated_keywords}

   Translated Extractive Summary:
   {translated_ex_summary}

   Translated Abstractive Summary:
   {translated_ab_summary}
   """  

    sections = []
    if full_text.strip():
       sections.append(f"Original Text: {full_text}")
    if extracted:
       sections.append(f"Extracted Details: {extracted}")
    if ex_summary.strip():
       sections.append(f"Extractive Summary: {ex_summary}")

    if ab_summary.strip():
        sections.append(f"Abstractive Summary: {ab_summary}")

    if translated_text.strip() and target_lang != "en":
        sections.append(f"Text ({target_lang}): {translated_text}")

    if translated_keywords:
        sections.append(f"Translated Extracted Details: {translated_keywords}")

    if translated_ex_summary.strip():
        sections.append(f"Translated Extractive Summary: {translated_ex_summary}")

    if translated_ab_summary.strip():
       sections.append(f"Translated Abstractive Summary: {translated_ab_summary}")

    full_text_for_voice = "\n\n".join(sections)
    audio_file = None
    if full_text_for_voice.strip():
        audio_file = speak_text_dynamic(full_text_for_voice, lang=target_lang)





# Generate voice (save as static file)
    audio_extracted = speak_text_dynamic(extracted, lang="en") if extracted else None
    audio_ex_summary = speak_text_dynamic(ex_summary, lang="en") if ex_summary else None
    audio_ab_summary = speak_text_dynamic(ab_summary, lang="en") if ab_summary else None
    audio_translation = speak_text_dynamic(translated_text, lang=target_lang) if translated_text else None

    audio_file = speak_text_dynamic(full_text_for_voice, lang=target_lang)


    return render_template(
        "index.html",
        extracted=extracted,
        detected_field=detected_field,
        full_text=full_text,
        translated_text=translated_text,
        translated_keywords=translated_keywords,
        target_lang=target_lang,
        ex_summary=ex_summary,
        ab_summary=ab_summary,
        translated_ex_summary=translated_ex_summary,
        translated_ab_summary=translated_ab_summary,
        files_by_category=files_by_category,
        audio_extracted=audio_extracted,
        audio_ex_summary=audio_ex_summary,
        audio_ab_summary=audio_ab_summary,
        audio_translation=audio_translation,
        audio_file=audio_file 
    ) '''
@app.route("/", methods=["GET", "POST"])
def index():
    extracted = None
    detected_field = None
    full_text = ""
    translated_text = ""
    translated_keywords = {}
    target_lang = "en"
    ex_summary = ""
    ab_summary = ""
    translated_ex_summary = ""
    translated_ab_summary = ""
    audio_combined = None

    if request.method == "POST":
        if "document" in request.files:
            file = request.files["document"]
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                temp_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(temp_path)

                # --- Extract text & keywords ---
                full_text = extract_text(temp_path)
                extracted = extract_keywords(full_text)
                detected_field = detect_document_field(full_text, extracted)

                # --- Move file into detected field folder ---
                field_folder = os.path.join(app.config["UPLOAD_FOLDER"], detected_field)
                os.makedirs(field_folder, exist_ok=True)
                new_path = os.path.join(field_folder, filename)
                os.replace(temp_path, new_path)

                # --- Summaries ---
                if full_text.strip():
                    ex_summary = extractive_summary(full_text)
                    ab_summary = abstractive_summary(full_text)

                # --- Translation ---
                target_lang = request.form.get("lang", "en")
                translated_text = full_text
                translated_keywords = extracted
                translated_ex_summary = ex_summary
                translated_ab_summary = ab_summary

                if target_lang != "en":
                    try:
                        translated_text = GoogleTranslator(source="auto", target=target_lang).translate(full_text)
                        translated_keywords = {
                            k: [GoogleTranslator(source="auto", target=target_lang).translate(v) for v in vals]
                            for k, vals in extracted.items()
                        }
                        translated_ex_summary = GoogleTranslator(source="auto", target=target_lang).translate(ex_summary)
                        translated_ab_summary = GoogleTranslator(source="auto", target=target_lang).translate(ab_summary)
                    except Exception as e:
                        translated_text = f"Translation error: {e}"
                        translated_keywords = extracted
                        translated_ex_summary = f"Translation error: {e}"
                        translated_ab_summary = f"Translation error: {e}"

                # --- Build full narration for TTS ---
                full_text_for_voice = f"""
Original Text:
{full_text}

Extracted Details:
{extracted}

Extractive Summary:
{ex_summary}

Abstractive Summary:
{ab_summary}

Text ({target_lang}):
{translated_text}

Translated Extracted Details:
{translated_keywords}

Translated Extractive Summary:
{translated_ex_summary}

Translated Abstractive Summary:
{translated_ab_summary}
"""
                # --- Generate audio only once ---
                if full_text_for_voice.strip():
                    try:
                        audio_combined = speak_text_dynamic(full_text_for_voice, lang=target_lang, prefix="combined")
                    except Exception as e:
                        print(f"TTS Error: {e}")
                        audio_combined = None

    # --- Always rescan uploads ---
    files_by_category = {}
    uploads_dir = app.config["UPLOAD_FOLDER"]
    for root, dirs, filenames in os.walk(uploads_dir):
        rel_dir = os.path.relpath(root, uploads_dir)
        if rel_dir == ".":
            rel_dir = "Uncategorized"
        files_by_category.setdefault(rel_dir, [])
        for f in filenames:
            rel_path = os.path.relpath(os.path.join(root, f), uploads_dir).replace("\\", "/")
            files_by_category[rel_dir].append(rel_path)

    return render_template(
        "index.html",
        extracted=extracted,
        detected_field=detected_field,
        full_text=full_text,
        translated_text=translated_text,
        translated_keywords=translated_keywords,
        target_lang=target_lang,
        ex_summary=ex_summary,
        ab_summary=ab_summary,
        translated_ex_summary=translated_ex_summary,
        translated_ab_summary=translated_ab_summary,
        files_by_category=files_by_category,
        audio_combined=audio_combined
    )


#from flask import send_from_directory


import os
from flask import Flask, render_template, request, url_for, send_from_directory, Response, abort



# Use absolute path for uploads
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    filename = filename.replace("\\", "/")
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    # ✅ Check if file exists (avoids 500 error)
    if not os.path.isfile(file_path):
        return abort(404, description=f"File not found: {filename}")

    # ✅ Handle .txt files (show in browser)
    if filename.lower().endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return Response(content, mimetype="text/plain")

    # ✅ Handle other files (pdf, docx, etc.)
    # If file is in a subfolder (education/health/etc.)
    #folder = os.path.dirname(filename)
    #fname = os.path.basename(filename)
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


if __name__ == "__main__":
   app.run(debug=True)